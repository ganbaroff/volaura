"""Tests for CV text extraction (Sprint 1)."""

import pytest
from app.services.cv_parser import extract_text, extract_text_from_docx, extract_text_from_pdf


class TestExtractText:
    """Tests for the extract_text dispatcher."""

    def test_unsupported_format_raises(self):
        with pytest.raises(ValueError, match="Unsupported file format"):
            extract_text(b"fake content", "resume.txt")

    def test_unsupported_exe_raises(self):
        with pytest.raises(ValueError, match="Unsupported file format"):
            extract_text(b"\x4d\x5a", "malware.exe")

    def test_empty_pdf_raises(self):
        """A valid but empty PDF should raise ValueError."""
        # Minimal valid PDF with no text content
        import pymupdf
        doc = pymupdf.open()
        doc.new_page()  # blank page
        pdf_bytes = doc.tobytes()
        doc.close()
        with pytest.raises(ValueError, match="no extractable text"):
            extract_text_from_pdf(pdf_bytes)

    def test_valid_pdf_extracts_text(self):
        """A PDF with text should return that text."""
        import pymupdf
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "John Doe\nSoftware Engineer\nPython, React, AWS")
        pdf_bytes = doc.tobytes()
        doc.close()

        result = extract_text(pdf_bytes, "resume.pdf")
        assert "John Doe" in result
        assert "Software Engineer" in result
        assert "Python" in result

    def test_valid_docx_extracts_text(self):
        """A DOCX with paragraphs should return text."""
        import docx
        import io
        doc = docx.Document()
        doc.add_paragraph("Jane Smith")
        doc.add_paragraph("Project Manager at COP29")
        doc.add_paragraph("ClickUp, Jira, Asana")
        buf = io.BytesIO()
        doc.save(buf)

        result = extract_text(buf.getvalue(), "cv.docx")
        assert "Jane Smith" in result
        assert "COP29" in result
        assert "ClickUp" in result

    def test_pdf_extension_case_insensitive(self):
        """Should handle .PDF uppercase."""
        import pymupdf
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Test Content")
        pdf_bytes = doc.tobytes()
        doc.close()

        result = extract_text(pdf_bytes, "RESUME.PDF")
        assert "Test Content" in result
